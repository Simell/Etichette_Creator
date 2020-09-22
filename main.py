#!/bin/python3


import sqlite3
from sqlite3 import Error
import os
import docx 
from docx import Document
from docx.shared import Pt
import subprocess
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

#debug 
print("Programma Partito")

def create_connection(db_file): 
    
    conn = None
    try:
        conn = sqlite3.connect(db_file)
        print(sqlite3.version)
    except Error as e:
        print(e)

    finally:
        if conn:
            conn.close()


def crea_table():
    try:  
        c.execute("""CREATE TABLE indirizzi 
                (file)""")
        conn.commit()
        print("tabella creata")
    except Error as e:
        
        print(e)


def inserimento_dati():
   #print( os.listdir('Indirizzi_Spedizioni'))
    for i in os.listdir('Indirizzi_SpedizioniDocx'):
        c.execute("INSERT INTO indirizzi (file) values(?)", (i,))
        conn.commit()
    print("dati inseriti")




# tutte le variabili 
conn = sqlite3.connect(r"indirizzi.db")
c = conn.cursor()

#crea_table()
#inserimento_dati()


# inizio comandi interfaccia utente 


# programma per creare nuova etichetta in formato docx 
def nuova_etichetta():
    # comandi
    nome = input("inserisci nome e cognome: ")
    azienda = input("inserisci azienda: ")
    indirizzo = input("inserisci indirizzo: ")
    località = input("inserisci località: ")   
    pagamento = input(" inserisci modalità di pagamento: ")
    prezzo = input("inserisci prezzo: ")
    prezzo_scritto = input("inserisci prezzo scritto: ")

    # modifica del documento
    try:
        document = Document("template/template.docx")
    except:
        os.mkdirs("template/")
        document = Document("template/template.docx")
    # creazione stili
    styles = document.styles
    
    # stile Times New Roman: solo bold con dimensione 26
    style = styles.add_style('Times New Roman', WD_STYLE_TYPE.PARAGRAPH)
    
    style.font.bold = True
    style.font.size = Pt(26)
    
    # stile Underline:  bold e sottolineato con dimensione 26
    style1 = styles.add_style('Underline', WD_STYLE_TYPE.PARAGRAPH) 
    style1.font.underline = True
    style1.font.bold = True
    style1.font.size = Pt(26)

    # formattazione documento
    par_nome = document.add_paragraph(nome)
    par_nome.alignment = 1

    par_azienda = document.add_paragraph(azienda)
    par_azienda.alignment = 1
    
    par_vuoto = document.add_paragraph(" ")    
    
    par_indirizzo = document.add_paragraph(indirizzo)
    par_indirizzo.alignment = 1
    
    par_vuoto = document.add_paragraph(" ")
    
    par_località = document.add_paragraph(località)
    par_località.alignment = 1
    
    par_vuoto = document.add_paragraph(" ")
    
    par_pagamento = document.add_paragraph(pagamento+" € "+prezzo)
    par_pagamento.alignment = 1
    par_prezzo_scritto = document.add_paragraph("("+prezzo_scritto+")")
    par_prezzo_scritto.alignment = 1

    # applicazione dello stile 
    par_nome.style = styles["Times New Roman"]
    par_azienda.style = styles["Times New Roman"]
    par_indirizzo.style = styles["Times New Roman"]

    par_località.style = styles["Underline"]
    par_pagamento.style = styles["Underline"]
    par_prezzo_scritto.style  = styles["Times New Roman"]

    
    # salvataggio documento
    document.save("Indirizzi_SpedizioniDocx/"+nome+".docx")



# programma di STAMPA
def stampa():
    file_stampa = input("Che file vuoi stampare?:")
    print("stampo il file:", file_stampa)
    try:
        subprocess.call(["C:\Windows\System32\WindowsPowerShell\\v1.0\powershell.exe", "soffice -p   '.\Indirizzi_SpedizioniDocx\{}'".format(file_stampa)])
    except: 
       print("XXXERROREXXX  file non trovato")

# definizione comando AIUTO
def aiuto():
    print("######AIUTO######")
    print("Scrivi 'stampa' per procedere con la stampa di un file")
    print("Scrivi 'nuova' per creare un nuovo documento")
    print("#################")




# interfaccia principale
def interfaccia():
    comando = input("-->")
    
    if comando == "":
        interfaccia()


    if comando == "stampa":
        stampa()
        interfaccia()
    
    if comando == "nuova":
        nuova_etichetta()
        interfaccia()
    
    if comando == "aiuto":
        aiuto()
        interfaccia()
    


    else:
        print("comando non trovato")
        interfaccia()

interfaccia()
